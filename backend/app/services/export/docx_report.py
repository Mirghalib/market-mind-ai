"""Professional .docx report builder for marketing strategies.

Produces a polished, editable Word document: cover page, table of
contents, colored section headings, styled tables, bullet/numbered
lists, a footer with page numbers, and clean typography. Uses only the
public python-docx API (no raw OOXML hacks) so the document opens and
edits cleanly in Microsoft Word.

The builder is imported lazily by the DocxRenderer so the renderer
registry stays importable without python-docx.
"""
from __future__ import annotations

from io import BytesIO

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from app.services.export.report_data import APP_NAME, BRAND, ReportData


def _rgb(hex_color: str):
    return RGBColor.from_string(hex_color.lstrip("#"))


def _shade_cell(cell, hex_color: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color.lstrip("#"))
    tc_pr.append(shd)


def _add_footer_page_number(document: Document, footer_text: str = "Market Mind AI") -> None:
    """Add a footer with the report name, page number and total pages."""
    section = document.sections[0]
    footer = section.footer
    footer.is_linked_to_previous = False
    paragraph = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    paragraph.text = ""
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    def _run(text: str, bold: bool = False, size: int = 8, color: str = BRAND["muted"]):
        run = paragraph.add_run(text)
        run.font.size = Pt(size)
        run.font.bold = bold
        run.font.color.rgb = _rgb(color)
        return run

    _run(footer_text, size=8)
    _run("   ·   ", size=8)

    # PAGE field
    page = OxmlElement("w:fldSimple")
    page.set(qn("w:instr"), "PAGE")
    page_run = OxmlElement("w:r")
    page_text = OxmlElement("w:t")
    page_text.text = "1"
    page_run.append(page_text)
    page.append(page_run)
    paragraph._p.append(page)

    _run(" of ", size=8)

    # NUMPAGES field
    total = OxmlElement("w:fldSimple")
    total.set(qn("w:instr"), "NUMPAGES")
    total_run = OxmlElement("w:r")
    total_text = OxmlElement("w:t")
    total_text.text = "1"
    total_run.append(total_text)
    total.append(total_run)
    paragraph._p.append(total)


def _style_base(document: Document) -> None:
    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = _rgb(BRAND["dark"])
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25


def _add_heading(document: Document, text: str, level: int = 1, color: str = BRAND["primary_dark"]) -> None:
    heading = document.add_heading(text, level=level)
    for run in heading.runs:
        run.font.color.rgb = _rgb(color)
        run.font.name = "Calibri"
    heading.paragraph_format.space_before = Pt(14 if level == 1 else 10)
    heading.paragraph_format.space_after = Pt(6)
    return heading


def _add_cover(document: Document, data: ReportData) -> None:
    # Top accent bar via a full-width table.
    bar = document.add_table(rows=1, cols=1)
    bar.autofit = True
    cell = bar.cell(0, 0)
    _shade_cell(cell, BRAND["primary"])
    cell.width = Inches(6.5)

    document.add_paragraph()
    document.add_paragraph()

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(data.name)
    run.font.size = Pt(30)
    run.font.bold = True
    run.font.color.rgb = _rgb(BRAND["slate"])

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("Marketing Strategy Report")
    run.font.size = Pt(15)
    run.font.color.rgb = _rgb(BRAND["primary"])

    document.add_paragraph()

    meta_lines = [
        ("Industry", data.industry),
        ("Country", data.country),
        ("Target audience", data.audience or "—"),
    ]
    if data.budget_label and data.budget_label != "Not specified":
        meta_lines.append(("Budget", data.budget_label))
    meta_lines.append(("Generated", data.generated_date or "—"))
    meta_lines.append(("Prepared by", APP_NAME))
    meta = document.add_table(rows=len(meta_lines), cols=2)
    meta.style = "Table Grid"
    for i, (label, value) in enumerate(meta_lines):
        label_cell, value_cell = meta.rows[i].cells
        label_cell.text = label.upper()
        for r in label_cell.paragraphs[0].runs:
            r.font.bold = True
            r.font.size = Pt(9)
            r.font.color.rgb = _rgb(BRAND["primary"])
        _shade_cell(label_cell, BRAND["light"])
        value_cell.text = value
        for r in value_cell.paragraphs[0].runs:
            r.font.size = Pt(10.5)
            r.font.color.rgb = _rgb(BRAND["dark"])

    document.add_paragraph()
    tagline = document.add_paragraph()
    tagline.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = tagline.add_run(f"{APP_NAME} — AI-Powered Marketing Intelligence")
    run.font.size = Pt(9)
    run.font.italic = True
    run.font.color.rgb = _rgb(BRAND["muted"])

    document.add_page_break()


def _add_toc(document: Document, sections: list[str]) -> None:
    _add_heading(document, "Table of Contents")
    for section in sections:
        p = document.add_paragraph()
        p.paragraph_format.space_after = Pt(3)
        run = p.add_run("•  " + section.replace("_", " ").title())
        run.font.size = Pt(11)
        run.font.color.rgb = _rgb(BRAND["dark"])
        run.font.bold = False
    document.add_page_break()


def _add_styled_table(document: Document, header: list[str], rows: list[list[str]],
                      col_widths: list[float] | None = None) -> None:
    table = document.add_table(rows=1 + len(rows), cols=len(header))
    table.style = "Table Grid"
    table.autofit = True

    # Header row.
    for j, label in enumerate(header):
        cell = table.cell(0, j)
        cell.text = label
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        _shade_cell(cell, BRAND["primary"])
        for r in cell.paragraphs[0].runs:
            r.font.bold = True
            r.font.size = Pt(9.5)
            r.font.color.rgb = _rgb(BRAND["white"])

    # Body rows.
    for i, row in enumerate(rows, start=1):
        for j, value in enumerate(row):
            cell = table.cell(i, j)
            cell.text = str(value)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            for r in cell.paragraphs[0].runs:
                r.font.size = Pt(9.5)
                r.font.color.rgb = _rgb(BRAND["dark"])
            if i % 2 == 0:
                _shade_cell(cell, BRAND["panel"])

    if col_widths:
        for j, width in enumerate(col_widths):
            for row in table.rows:
                row.cells[j].width = Inches(width)


def _add_bullets(document: Document, items: list[str], style: str = "List Bullet") -> None:
    for item in items:
        p = document.add_paragraph(style=style)
        run = p.add_run(item)
        run.font.size = Pt(10.5)
        run.font.color.rgb = _rgb(BRAND["dark"])


def _add_numbered(document: Document, items: list[str]) -> None:
    for item in items:
        p = document.add_paragraph(style="List Number")
        run = p.add_run(item)
        run.font.size = Pt(10.5)
        run.font.color.rgb = _rgb(BRAND["dark"])


def _add_kpi_band(document: Document, data: ReportData) -> None:
    """Purple KPI dashboard row (score, objectives, channels)."""
    stats = [
        ("Marketing Score", f"{data.score}/100"),
        ("Objectives", str(len(data.objectives))),
        ("Channels", str(len(data.channels))),
    ]
    table = document.add_table(rows=1, cols=len(stats))
    for i, (label, value) in enumerate(stats):
        cell = table.cell(0, i)
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(value)
        run.font.size = Pt(18)
        run.font.bold = True
        run.font.color.rgb = _rgb(BRAND["primary"])
        p2 = cell.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run2 = p2.add_run(label.upper())
        run2.font.size = Pt(8)
        run2.font.color.rgb = _rgb(BRAND["muted"])
        _shade_cell(cell, BRAND["light"])


def _section_executive_summary(document: Document, data: ReportData) -> None:
    _add_heading(document, "Executive Summary")
    if data.summary_text:
        p = document.add_paragraph()
        run = p.add_run(data.summary_text)
        run.font.size = Pt(11)
    if data.highlights:
        _add_heading(document, "Key Highlights", level=2, color=BRAND["secondary"])
        _add_bullets(document, data.highlights)
    if data.recommendation:
        p = document.add_paragraph()
        run = p.add_run(f"Recommendation: {data.recommendation}")
        run.font.bold = True
        run.font.size = Pt(10.5)
    _add_kpi_band(document, data)


def _section_score(document: Document, data: ReportData) -> None:
    _add_heading(document, "Marketing Score")
    if data.score_summary:
        p = document.add_paragraph()
        run = p.add_run(data.score_summary)
        run.font.size = Pt(11)
    if data.score_benchmark:
        p = document.add_paragraph()
        run = p.add_run(f"Benchmark: {data.score_benchmark}")
        run.font.italic = True
        run.font.size = Pt(9.5)
        run.font.color.rgb = _rgb(BRAND["muted"])
    if data.score_breakdown:
        rows = [[b["area"], f"{b['score']}/100", b["assessment"]] for b in data.score_breakdown]
        _add_styled_table(document, ["Area", "Score", "Assessment"], rows,
                          col_widths=[1.6, 0.9, 4.0])


def _section_market(document: Document, data: ReportData) -> None:
    _add_heading(document, "Market Overview")
    mkt = data.market
    if mkt.get("summary"):
        p = document.add_paragraph()
        run = p.add_run(str(mkt["summary"]))
        run.font.size = Pt(11)
    if data.market_trends:
        _add_heading(document, "Market Trends", level=2, color=BRAND["secondary"])
        _add_bullets(document, data.market_trends)
    if data.market_drivers:
        _add_heading(document, "Key Drivers", level=2, color=BRAND["secondary"])
        _add_bullets(document, data.market_drivers)


def _section_persona(document: Document, data: ReportData) -> None:
    _add_heading(document, "Customer Persona")
    persona = data.persona
    rows = [
        ["Name", clean(persona.get("name"))],
        ["Age range", clean(persona.get("ageRange"))],
        ["Location", clean(persona.get("location"))],
        ["Occupation", clean(persona.get("occupation"))],
        ["Income level", clean(persona.get("incomeLevel"))],
    ]
    _add_styled_table(document, ["Attribute", "Profile"], rows, col_widths=[1.8, 4.7])
    for label, key in [("Interests", "interests"), ("Pain Points", "painPoints"),
                       ("Goals", "goals"), ("Preferred Channels", "preferredChannels")]:
        items = [clean(i) for i in (persona.get(key) or []) if clean(i)]
        if items:
            _add_heading(document, label, level=2, color=BRAND["secondary"])
            _add_bullets(document, items)


def _section_swot(document: Document, data: ReportData) -> None:
    _add_heading(document, "SWOT Analysis")
    quadrants = data.swot_quadrants
    table = document.add_table(rows=2, cols=2)
    table.style = "Table Grid"
    labels = [("Strengths", "Weaknesses"), ("Opportunities", "Threats")]
    for i in range(2):
        for j in range(2):
            cell = table.cell(i, j)
            label = labels[i][j]
            cell.text = label
            _shade_cell(cell, BRAND["primary"] if j == 0 else BRAND["secondary"])
            for r in cell.paragraphs[0].runs:
                r.font.bold = True
                r.font.color.rgb = _rgb(BRAND["white"])
            for item in quadrants[label]:
                p = cell.add_paragraph()
                run = p.add_run(f"• {item}")
                run.font.size = Pt(9.5)
    if data.swot.get("overallAssessment"):
        p = document.add_paragraph()
        run = p.add_run(str(data.swot["overallAssessment"]))
        run.font.italic = True
        run.font.size = Pt(10.5)


def _section_strategy(document: Document, data: ReportData) -> None:
    _add_heading(document, "Marketing Strategy")
    if data.objectives:
        _add_heading(document, "Objectives", level=2, color=BRAND["secondary"])
        _add_numbered(document, data.objectives)
    if data.channels:
        rows = [[c["name"], c["priority"].capitalize(), c["description"]]
                for c in data.channels]
        _add_styled_table(document, ["Channel", "Priority", "Why / How"], rows,
                          col_widths=[1.6, 1.2, 3.7])
    if data.budget:
        _add_heading(document, "Budget Allocation", level=2, color=BRAND["secondary"])
        _add_styled_table(document, ["Channel", "Allocation"], data.budget_rows,
                          col_widths=[4.0, 1.5])


def _section_kpis(document: Document, data: ReportData) -> None:
    if not data.kpis:
        return
    _add_heading(document, "Key Performance Indicators")
    _add_styled_table(document, ["Metric", "Target", "Timeframe"], data.kpi_rows,
                      col_widths=[2.6, 1.8, 1.6])


def _section_competitors(document: Document, data: ReportData) -> None:
    if not data.competitors:
        return
    _add_heading(document, "Competitor Analysis")
    rows = [[c["name"], c["position"], c["threat"].capitalize(),
             ", ".join(str(s) for s in c["strengths"]),
             ", ".join(str(w) for w in c["weaknesses"])] for c in data.competitors]
    _add_styled_table(document, ["Competitor", "Position", "Threat", "Strengths", "Weaknesses"], rows)


def _section_roi(document: Document, data: ReportData) -> None:
    if not data.roi_projections:
        return
    _add_heading(document, "Estimated ROI")
    if data.roi.get("summary"):
        p = document.add_paragraph()
        run = p.add_run(str(data.roi["summary"]))
        run.font.size = Pt(11)
    _add_styled_table(document, ["Period", "Investment", "Projected Return", "ROI"], data.roi_rows,
                      col_widths=[1.5, 1.5, 1.8, 1.0])
    if data.roi.get("paybackPeriod"):
        p = document.add_paragraph()
        run = p.add_run(f"Payback period: {data.roi['paybackPeriod']}")
        run.font.bold = True
        run.font.size = Pt(10.5)


def _section_roadmap(document: Document, data: ReportData) -> None:
    if not data.roadmap_phases:
        return
    _add_heading(document, "90-Day Implementation Roadmap")
    for phase in data.roadmap_phases:
        name = clean(phase.get("name"))
        duration = clean(phase.get("duration"))
        if name:
            p = document.add_paragraph()
            run = p.add_run(name + (f" — {duration}" if duration else ""))
            run.font.bold = True
            run.font.size = Pt(11)
            run.font.color.rgb = _rgb(BRAND["primary_dark"])
        for label, key in [("Objectives", "objectives"),
                           ("Key Activities", "keyActivities"),
                           ("Success Metrics", "successMetrics")]:
            items = [clean(i) for i in (phase.get(key) or []) if clean(i)]
            if items:
                _add_heading(document, label, level=2, color=BRAND["secondary"])
                _add_bullets(document, items)


def _section_recommendations(document: Document, data: ReportData) -> None:
    rec = data.recommendations
    if not rec:
        return
    _add_heading(document, "Final Recommendations")
    if rec.get("summary"):
        p = document.add_paragraph()
        run = p.add_run(str(rec["summary"]))
        run.font.size = Pt(11)
    for label, key in [("Top Priorities", "priorities"),
                       ("Quick Wins", "quickWins"),
                       ("Long-Term Investments", "longTermInvestments"),
                       ("Success Criteria", "successCriteria")]:
        items = [clean(i) for i in (rec.get(key) or []) if clean(i)]
        if items:
            _add_heading(document, label, level=2, color=BRAND["secondary"])
            _add_bullets(document, items)
    if rec.get("closingStatement"):
        p = document.add_paragraph()
        run = p.add_run(str(rec["closingStatement"]))
        run.font.italic = True
        run.font.size = Pt(10.5)


def clean(value) -> str:
    return str(value) if value is not None else "—"


def build_strategy_report_docx(strategy) -> bytes:
    """Render the full professional report as .docx bytes."""
    data = ReportData(strategy)
    buffer = BytesIO()
    document = Document()

    _style_base(document)
    # Narrow-ish margins for a consulting look.
    for section in document.sections:
        section.left_margin = Inches(0.9)
        section.right_margin = Inches(0.9)
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)

    _add_cover(document, data)

    # TOC from the sections present in the content.
    toc_names = {
        "executiveSummary": "Executive Summary",
        "marketingScore": "Marketing Score",
        "marketOverview": "Market Overview",
        "customerPersona": "Customer Persona",
        "swotAnalysis": "SWOT Analysis",
        "competitorAnalysis": "Competitor Analysis",
        "marketingStrategy": "Marketing Strategy",
        "seoKeywords": "SEO Strategy",
        "emailCampaign": "Email Marketing",
        "socialMediaStrategy": "Social Media Strategy",
        "advertisementIdeas": "Paid Advertising",
        "contentCalendar": "Content Calendar",
        "implementationRoadmap": "90-Day Roadmap",
        "weeklyMilestones": "Weekly Milestones",
        "estimatedROI": "Estimated ROI",
        "riskMitigation": "Risks & Mitigation",
        "finalRecommendations": "Final Recommendations",
    }
    present = [toc_names.get(k, k.replace("_", " ").title()) for k in data.present_sections]
    _add_toc(document, present)

    # Section content.
    _section_executive_summary(document, data)
    _section_score(document, data)
    _section_market(document, data)
    _section_persona(document, data)
    _section_swot(document, data)
    _section_strategy(document, data)
    _section_kpis(document, data)
    _section_competitors(document, data)
    _section_roi(document, data)
    _section_roadmap(document, data)
    _section_recommendations(document, data)

    _add_footer_page_number(document, f"Market Mind AI — {data.name}")

    document.save(buffer)
    return buffer.getvalue()
